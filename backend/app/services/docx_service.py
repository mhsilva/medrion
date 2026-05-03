"""
DOCX generation service for Medrion prescriptions.
Produces A4-formatted .docx files with proper styling.
"""

import io
from datetime import date
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _set_cell_background(cell, fill_color: str) -> None:
    """Set a table cell background color using OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def _set_left_border(cell, color: str = "FF0000", width: int = 24) -> None:
    """Add a colored left border to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color)
    tcBorders.append(left)
    tcPr.append(tcBorders)


def _add_alert_paragraph(doc: Document, text: str) -> None:
    """Add a specially styled alert paragraph using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_background(cell, "FFF0F0")
    _set_left_border(cell, "CC0000", 36)
    para = cell.paragraphs[0]
    para.paragraph_format.left_indent = Cm(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True
    doc.add_paragraph()  # spacer


def _add_section_paragraph(doc: Document, text: str) -> None:
    """Add a regular prescription body paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def generate_docx(
    prescription_text: str, doctor_header: dict[str, Any], patient: dict[str, Any]
) -> bytes:
    """
    Generate a .docx file for the prescription.

    Args:
        prescription_text: The full prescription text (11 sections).
        doctor_header: Dict with doctor info (name, crm, state, specialty, address, phone, email).
        patient: Dict with patient info (name, birth_date, age).

    Returns:
        Raw bytes of the .docx file.
    """
    doc = Document()

    # Page setup: A4, 2cm margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ------------------------------------------------------------------ #
    # Header block                                                         #
    # ------------------------------------------------------------------ #
    def _header_run(para, text: str, bold: bool = False, size: int = 12) -> None:
        run = para.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _header_run(name_para, doctor_header.get("name", ""), bold=True, size=14)

    crm_line = []
    if doctor_header.get("crm"):
        crm_line.append(f"CRM {doctor_header['crm']}")
    if doctor_header.get("state"):
        crm_line[-1] += f"/{doctor_header['state']}" if crm_line else doctor_header["state"]
    if doctor_header.get("specialty"):
        crm_line.append(doctor_header["specialty"])

    if crm_line:
        crm_para = doc.add_paragraph()
        crm_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _header_run(crm_para, " — ".join(crm_line), size=11)

    contact_parts = []
    if doctor_header.get("address"):
        contact_parts.append(doctor_header["address"])
    if doctor_header.get("phone"):
        contact_parts.append(f"Tel: {doctor_header['phone']}")
    if doctor_header.get("email"):
        contact_parts.append(doctor_header["email"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _header_run(contact_para, " | ".join(contact_parts), size=10)

    # Separator line
    sep_para = doc.add_paragraph()
    sep_run = sep_para.add_run("─" * 80)
    sep_run.font.name = "Calibri"
    sep_run.font.size = Pt(9)
    sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ------------------------------------------------------------------ #
    # Patient identification block                                         #
    # ------------------------------------------------------------------ #
    patient_name = patient.get("name", "")
    birth_date = patient.get("birth_date", "")
    patient_age = patient.get("age", "")
    prescription_date = date.today().strftime("%d/%m/%Y")

    patient_para = doc.add_paragraph()
    _header_run(patient_para, "Paciente: ", bold=True, size=11)
    _header_run(patient_para, patient_name, size=11)

    details_para = doc.add_paragraph()
    detail_parts = []
    if birth_date:
        detail_parts.append(f"Data de nascimento: {birth_date}")
    if patient_age:
        detail_parts.append(f"Idade: {patient_age} anos")
    detail_parts.append(f"Data: {prescription_date}")
    _header_run(details_para, "  |  ".join(detail_parts), size=10)

    sep_para2 = doc.add_paragraph()
    sep_run2 = sep_para2.add_run("─" * 80)
    sep_run2.font.name = "Calibri"
    sep_run2.font.size = Pt(9)
    sep_run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # spacer

    # ------------------------------------------------------------------ #
    # Prescription body — parse [ALERTA] blocks                           #
    # ------------------------------------------------------------------ #
    lines = prescription_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if stripped.upper().startswith("[ALERTA]"):
            _add_alert_paragraph(doc, stripped)
        elif stripped == "":
            doc.add_paragraph()
        else:
            _add_section_paragraph(doc, stripped)

    # ------------------------------------------------------------------ #
    # Signature block                                                      #
    # ------------------------------------------------------------------ #
    doc.add_paragraph()
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _header_run(sig_para, "_" * 40, size=11)

    sig_name_para = doc.add_paragraph()
    sig_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _header_run(sig_name_para, doctor_header.get("name", ""), bold=True, size=11)

    crm_sig_parts = []
    if doctor_header.get("crm"):
        crm_sig_parts.append(f"CRM {doctor_header['crm']}")
    if doctor_header.get("state"):
        crm_sig_parts[-1] += f"/{doctor_header['state']}" if crm_sig_parts else doctor_header["state"]
    if crm_sig_parts:
        sig_crm_para = doc.add_paragraph()
        sig_crm_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _header_run(sig_crm_para, crm_sig_parts[0], size=10)

    # ------------------------------------------------------------------ #
    # Footer                                                               #
    # ------------------------------------------------------------------ #
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Documento gerado via Medrion — {prescription_date}"
    )
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
